"""Ingestion lineage: what each index attempt actually produced.

WO-2026-08-14-INGESTION-INSTRUMENTATION-001 -- AC-01 .. AC-06.

The centre of this file is ``TestParserOutputIsUnchanged``, which is the AC-03 proof. Its
expected values were produced by running the parser on commit 12f0216 -- BEFORE any line of
this Work Order's implementation existed -- and pasted in verbatim. The Work Order says the
existing suite passing is "necessary but not sufficient"; a suite can only catch a regression
in what it already asserts, and nothing asserted the DOCX/PDF chunk ids or the Korean-heading
markdown locators. So the pin is a digest over every field of every chunk for four inputs,
plus the individual locators spelled out so a reader can see what is pinned rather than
trusting a hash.
"""

from __future__ import annotations

import hashlib
import importlib.util
import io
import json
from collections.abc import Iterator

import pytest

RUNTIME_DEPS = ("fastapi", "httpx", "pydantic_settings", "sqlalchemy", "docx", "pypdf")


def runtime_deps_available() -> bool:
    return all(importlib.util.find_spec(package) for package in RUNTIME_DEPS)


pytestmark = pytest.mark.skipif(
    not runtime_deps_available(),
    reason="Runtime dependencies are not installed",
)


# =======================================================================================
# Fixture content: one policy text, expressed as markdown and as a DOCX with the same
# headings. This is the pair the Work Order's problem statement is about.
# =======================================================================================
POLICY_MD = (
    "# 취업규칙\n"
    "\n"
    "## 제3장 연차 유급휴가\n"
    "\n"
    "제15조 근속 1년 이상 직원은 15일의 연차 유급휴가를 받는다.\n"
    "연차는 회계연도 기준으로 산정한다.\n"
    "\n"
    "## 제4장 휴가 신청 절차\n"
    "\n"
    "제20조 휴가는 사용 3일 전까지 결재선에 상신한다.\n"
    "긴급한 경우 사후 결재를 허용한다.\n"
    "\n"
    "## 제5장 경조사 휴가\n"
    "\n"
    "제25조 본인 결혼은 5일, 배우자 출산은 10일의 경조사 휴가를 부여한다.\n"
)


def build_docx_bytes(markdown_text: str) -> bytes:
    """The same text as a DOCX: every markdown line becomes a paragraph, '#' stripped.

    Deliberately NOT using Word heading styles -- the point of the contrast is that the
    chunker is told ``text/plain`` regardless, so even a perfectly structured DOCX loses its
    hierarchy at the seam described in the ingestion-normalization design section 1-1.
    """
    from docx import Document as DocxDocument

    document = DocxDocument()
    for line in markdown_text.split("\n"):
        stripped = line.lstrip("#").strip()
        if stripped:
            document.add_paragraph(stripped)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(page_texts: list[str]) -> bytes:
    """A minimal multi-page PDF, hand-assembled so no new dependency is needed.

    Pages whose text is empty get no content stream at all, which is how a scanned page
    behaves to a text extractor: it is present and contributes nothing.
    """
    objects: list[bytes] = []

    def add(body: bytes) -> int:
        objects.append(body)
        return len(objects)

    font_id = add(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    page_ids: list[int] = []
    content_ids: list[int | None] = []
    for text in page_texts:
        if text:
            escaped = text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")
            stream = f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("latin-1", "replace")
            content_ids.append(
                add(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
            )
        else:
            content_ids.append(None)
        page_ids.append(add(b"PLACEHOLDER"))

    pages_id = add(b"PLACEHOLDER")
    catalog_id = add(b"<< /Type /Catalog /Pages " + str(pages_id).encode() + b" 0 R >>")

    for page_id, content_id in zip(page_ids, content_ids, strict=True):
        contents = f" /Contents {content_id} 0 R" if content_id else ""
        objects[page_id - 1] = (
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 792]"
            f" /Resources << /Font << /F1 {font_id} 0 R >> >>{contents} >>"
        ).encode()
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[pages_id - 1] = (
        f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>"
    ).encode()

    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets[1:]:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode()
    return bytes(out)


# =======================================================================================
# AC-03: no conversion, chunking or citation-locator output changes
# =======================================================================================
def _canonical(chunks) -> list[dict]:
    return [
        {
            "chunk_id": chunk.chunk_id,
            "chunk_index": chunk.chunk_index,
            "content": chunk.content,
            "content_hash": chunk.content_hash,
            "chunk_hash": chunk.chunk_hash,
            "token_count": chunk.token_count,
            "line_start": chunk.line_start,
            "line_end": chunk.line_end,
            "section_path": list(chunk.section_path),
            "citation_locator": chunk.citation_locator,
            "parser_version": chunk.parser_version,
            "chunker_version": chunk.chunker_version,
        }
        for chunk in chunks
    ]


def _parser_output_snapshot() -> dict:
    from app.domain.parsers import (
        DOCX_MIME_TYPE,
        chunker_mime_type_for,
        extract_text_from_bytes,
        parse_txt_md_document,
    )

    docx_text = extract_text_from_bytes(
        mime_type=DOCX_MIME_TYPE, content=build_docx_bytes(POLICY_MD)
    )
    return {
        "docx_extracted_text": docx_text,
        "docx": _canonical(
            parse_txt_md_document(
                document_id="doc-docx",
                document_version="2026-08-14",
                title="취업규칙 (DOCX)",
                mime_type=chunker_mime_type_for(DOCX_MIME_TYPE),
                source_text=docx_text,
            )
        ),
        "markdown": _canonical(
            parse_txt_md_document(
                document_id="doc-md",
                document_version="2026-08-14",
                title="취업규칙 (MD)",
                mime_type="text/markdown",
                source_text=POLICY_MD,
            )
        ),
        "plain": _canonical(
            parse_txt_md_document(
                document_id="doc-txt",
                document_version="v0",
                title="Holiday Policy",
                mime_type="text/plain",
                source_text="First paragraph.\n\nSecond paragraph for citation.",
            )
        ),
        "windowed": _canonical(
            parse_txt_md_document(
                document_id="doc-overlap",
                document_version="v0",
                title="Window Doc",
                mime_type="text/plain",
                source_text="\n".join(f"w{n}" for n in range(1, 13)),
                target_tokens=5,
                overlap_tokens=2,
            )
        ),
    }


#: sha256 of the canonical JSON of every field of every chunk for the four inputs above,
#: recorded from commit 12f0216 (the merge-base of this Work Order's branch), i.e. from the
#: tree BEFORE the lineage instrumentation existed. If instrumentation ever perturbs parsing
#: -- a normalisation "while we're here", a different extraction order, an extra strip -- this
#: digest changes and the Work Order's escalation clause applies: stop and report, because
#: behaviour change belongs to a separate Work Order.
PARSER_OUTPUT_DIGEST_BEFORE = (
    "702efdc09a5c945eb1b582c85a4dbd66b5792cdbc8610d5067711679b65af6d8"
)


class TestParserOutputIsUnchanged:
    """AC-03."""

    def test_every_parser_field_matches_the_pre_change_recording(self):
        snapshot = _parser_output_snapshot()
        digest = hashlib.sha256(
            json.dumps(snapshot, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode(
                "utf-8"
            )
        ).hexdigest()
        assert digest == PARSER_OUTPUT_DIGEST_BEFORE, (
            "Parser output changed. This Work Order is instrumentation only: conversion, "
            "chunking and citation locators must be byte-identical. Snapshot was: "
            f"{json.dumps(snapshot, ensure_ascii=False, sort_keys=True)}"
        )

    def test_the_pinned_values_are_the_ones_a_reader_would_check(self):
        """The digest above is opaque; these are the same facts in readable form."""
        snapshot = _parser_output_snapshot()
        assert [chunk["citation_locator"] for chunk in snapshot["markdown"]] == [
            "취업규칙 (MD) / 취업규칙 > 제3장 연차 유급휴가 / lines 5-6",
            "취업규칙 (MD) / 취업규칙 > 제4장 휴가 신청 절차 / lines 10-11",
            "취업규칙 (MD) / 취업규칙 > 제5장 경조사 휴가 / lines 15-15",
        ]
        assert [chunk["citation_locator"] for chunk in snapshot["docx"]] == [
            "취업규칙 (DOCX) / body / lines 1-17",
        ]
        assert [chunk["chunk_id"] for chunk in snapshot["docx"]] == [
            "doc-docx:2026-08-14:l1-17:c000:0d08caeb",
        ]
        assert [chunk["chunk_id"] for chunk in snapshot["markdown"]] == [
            "doc-md:2026-08-14:l5-6:c000:b651aa8b",
            "doc-md:2026-08-14:l10-11:c001:038f0d1d",
            "doc-md:2026-08-14:l15-15:c002:8bf4b5db",
        ]

    def test_the_collapse_this_work_order_measures_is_still_present(self):
        """Instrumentation must MEASURE the defect, not repair it. If a future slice fixes the
        seam, this test fails and that is correct -- it is the ledger entry saying the
        instrumentation baseline was taken against the broken behaviour."""
        snapshot = _parser_output_snapshot()
        assert all(chunk["section_path"] == [] for chunk in snapshot["docx"])
        assert all(chunk["section_path"] for chunk in snapshot["markdown"])


# =======================================================================================
# Fixtures
# =======================================================================================
@pytest.fixture
def client():
    from fastapi.testclient import TestClient
    from sqlalchemy import create_engine
    from sqlalchemy.orm import Session, sessionmaker
    from sqlalchemy.pool import StaticPool

    from app.core.database import Base, get_db
    from app.domain import models  # noqa: F401
    from app.main import create_app

    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)
    Base.metadata.create_all(bind=engine)

    def override_get_db() -> Iterator[Session]:
        db = testing_session()
        try:
            yield db
        finally:
            db.close()

    app = create_app()
    app.dependency_overrides[get_db] = override_get_db
    with TestClient(app) as test_client:
        test_client.lineage_session = testing_session
        yield test_client
        test_client.app.dependency_overrides.clear()
    Base.metadata.drop_all(bind=engine)


_KM = {"X-Agent-Forge-User": "km", "X-Agent-Forge-Roles": "knowledge-manager"}

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _create_source(client) -> dict:
    response = client.post(
        "/api/v1/knowledge/sources",
        headers=_KM,
        json={"name": "Policies", "description": "", "owner_department": "Operations"},
    )
    assert response.status_code == 201, response.text
    return response.json()


def _upload(client, source_id: str, *, title: str, filename: str, content: bytes, mime: str):
    return client.post(
        "/api/v1/knowledge/documents/upload",
        headers=_KM,
        data={
            "knowledge_source_id": source_id,
            "title": title,
            "confidentiality_level": "internal",
            "access_groups": "all-employees",
        },
        files={"file": (filename, content, mime)},
    )


def _lineage_rows(client, document_id: str | None = None) -> list:
    """Read document_ingestions straight from the table the API wrote to."""
    from sqlalchemy import select

    from app.domain.models import DocumentIngestion

    with client.lineage_session() as session:
        statement = select(DocumentIngestion).order_by(DocumentIngestion.created_at)
        if document_id is not None:
            statement = statement.where(DocumentIngestion.document_id == document_id)
        return list(session.scalars(statement))


def _locators(client, document_id: str) -> list[str]:
    response = client.get(f"/api/v1/knowledge/documents/{document_id}/chunks", headers=_KM)
    assert response.status_code == 200, response.text
    return [chunk["citation_locator"] for chunk in response.json()]


# =======================================================================================
# AC-01: markdown and DOCX with identical text produce visibly different lineage
# =======================================================================================
class TestStructureLossIsVisible:
    """AC-01. The two rows must differ in a way that IDENTIFIES the loss, not merely differ."""

    def _index_both(self, client):
        source = _create_source(client)
        markdown = _upload(
            client,
            source["id"],
            title="취업규칙 (MD)",
            filename="policy.md",
            content=POLICY_MD.encode("utf-8"),
            mime="text/markdown",
        )
        assert markdown.status_code == 201, markdown.text
        docx = _upload(
            client,
            source["id"],
            title="취업규칙 (DOCX)",
            filename="policy.docx",
            content=build_docx_bytes(POLICY_MD),
            mime=DOCX_MIME,
        )
        assert docx.status_code == 201, docx.text
        return markdown.json(), docx.json()

    def test_structured_chunk_count_differs_and_names_the_loss(self, client):
        from app.domain.ingestion_lineage import (
            WARNING_HEADING_DETECTION_UNAVAILABLE,
            WARNING_NO_HEADINGS_DETECTED,
        )

        markdown, docx = self._index_both(client)
        (md_row,) = _lineage_rows(client, markdown["document"]["id"])
        (docx_row,) = _lineage_rows(client, docx["document"]["id"])

        # The markdown path: every chunk carries a clause path.
        assert md_row.chunk_count == 3
        assert md_row.structured_chunk_count == 3
        assert md_row.converted_mime_type == "text/markdown"
        assert md_row.warnings == []

        # The binary path: ONE chunk, and it carries no clause path. Both effects recorded.
        assert docx_row.chunk_count == 1
        assert docx_row.structured_chunk_count == 0
        assert docx_row.converted_mime_type == "text/plain"
        assert set(docx_row.warnings) == {
            WARNING_HEADING_DETECTION_UNAVAILABLE,
            WARNING_NO_HEADINGS_DETECTED,
        }

    def test_the_warning_is_raised_for_the_binary_path_only(self, client):
        markdown, docx = self._index_both(client)
        (md_row,) = _lineage_rows(client, markdown["document"]["id"])
        (docx_row,) = _lineage_rows(client, docx["document"]["id"])
        assert md_row.warnings == []
        assert docx_row.warnings != []

    def test_the_lineage_matches_the_citations_the_reader_actually_gets(self, client):
        """The row is only worth anything if it corresponds to the citations produced."""
        markdown, docx = self._index_both(client)
        md_locators = _locators(client, markdown["document"]["id"])
        docx_locators = _locators(client, docx["document"]["id"])

        assert md_locators == [
            "취업규칙 (MD) / 취업규칙 > 제3장 연차 유급휴가 / lines 5-6",
            "취업규칙 (MD) / 취업규칙 > 제4장 휴가 신청 절차 / lines 10-11",
            "취업규칙 (MD) / 취업규칙 > 제5장 경조사 휴가 / lines 15-15",
        ]
        assert docx_locators == ["취업규칙 (DOCX) / body / lines 1-17"]

        (md_row,) = _lineage_rows(client, markdown["document"]["id"])
        (docx_row,) = _lineage_rows(client, docx["document"]["id"])
        assert md_row.structured_chunk_count == sum(
            1 for locator in md_locators if " / body / " not in locator
        )
        assert docx_row.structured_chunk_count == sum(
            1 for locator in docx_locators if " / body / " not in locator
        )

    def test_the_converter_and_its_version_are_recorded(self, client):
        markdown, docx = self._index_both(client)
        (md_row,) = _lineage_rows(client, markdown["document"]["id"])
        (docx_row,) = _lineage_rows(client, docx["document"]["id"])
        assert docx_row.converter_chain.startswith("python-docx/")
        assert docx_row.converter_chain.endswith(">text/plain")
        assert "unknown" not in docx_row.converter_chain
        assert md_row.converter_chain == "identity/decode-utf8>text/markdown"

    def test_the_declared_and_the_converted_mime_type_are_both_kept(self, client):
        """One column alone cannot express "the uploader sent DOCX, the chunker saw plain"."""
        _, docx = self._index_both(client)
        (docx_row,) = _lineage_rows(client, docx["document"]["id"])
        assert docx_row.source_mime_type == DOCX_MIME
        assert docx_row.converted_mime_type == "text/plain"


class TestHeadingAwareMimeTypesMatchTheParser:
    """The lineage module duplicates the parser's ``is_markdown`` set; prove they agree.

    Behavioural, not a constant comparison: it runs the real parser on a heading for every
    supported text MIME type and checks that "the parser produced a section path" is exactly
    "the lineage module considers this type heading-aware". If a future change makes
    ``text/plain`` heading-aware, this fails instead of the warning quietly becoming a lie.
    """

    def test_every_supported_text_mime_type_agrees(self):
        from app.domain.ingestion_lineage import HEADING_AWARE_MIME_TYPES
        from app.domain.parsers import SUPPORTED_TEXT_MIME_TYPES, parse_txt_md_document

        for mime_type in sorted(SUPPORTED_TEXT_MIME_TYPES):
            chunks = parse_txt_md_document(
                document_id="d",
                document_version="v0",
                title="T",
                mime_type=mime_type,
                source_text="# Heading\n\nbody text.",
            )
            parser_detects_headings = any(chunk.section_path for chunk in chunks)
            assert parser_detects_headings == (mime_type in HEADING_AWARE_MIME_TYPES), mime_type


# =======================================================================================
# AC-02: append-only across re-indexing
# =======================================================================================
class TestLineageIsAppendOnly:
    """AC-02. A re-index ADDS a row. This is the property a converter incident depends on."""

    def _register_and_index(self, client) -> str:
        source = _create_source(client)
        uploaded = _upload(
            client,
            source["id"],
            title="취업규칙 (MD)",
            filename="policy.md",
            content=POLICY_MD.encode("utf-8"),
            mime="text/markdown",
        )
        assert uploaded.status_code == 201, uploaded.text
        return uploaded.json()["document"]["id"]

    def _reindex(self, client, document_id: str, source_text: str):
        response = client.post(
            f"/api/v1/knowledge/documents/{document_id}/index-jobs",
            headers=_KM,
            json={"force_reindex": True, "source_text": source_text},
        )
        assert response.status_code == 201, response.text
        return response.json()

    def test_reindexing_leaves_two_rows_with_distinct_job_ids(self, client):
        document_id = self._register_and_index(client)
        self._reindex(client, document_id, POLICY_MD)

        rows = _lineage_rows(client, document_id)
        assert len(rows) == 2
        job_ids = [row.index_job_id for row in rows]
        assert all(job_ids)
        assert len(set(job_ids)) == 2

    def test_the_earlier_attempt_keeps_its_own_numbers(self, client):
        """The failure mode this criterion guards: a re-index that OVERWRITES, leaving no
        record that the earlier attempt produced something different."""
        document_id = self._register_and_index(client)
        # Re-index with text that has NO headings -- the second attempt collapses.
        self._reindex(client, document_id, "제15조 연차는 15일이다.\n제20조 3일 전에 신청한다.")

        first, second = _lineage_rows(client, document_id)
        assert (first.chunk_count, first.structured_chunk_count) == (3, 3)
        assert (second.chunk_count, second.structured_chunk_count) == (1, 0)
        assert first.warnings == []
        assert "NO_HEADINGS_DETECTED" in second.warnings
        assert first.index_job_id != second.index_job_id

    def test_a_third_attempt_adds_a_third_row(self, client):
        document_id = self._register_and_index(client)
        self._reindex(client, document_id, POLICY_MD)
        self._reindex(client, document_id, POLICY_MD)
        assert len(_lineage_rows(client, document_id)) == 3

    def test_the_pipeline_never_updates_an_existing_row(self):
        """Structural, so it cannot regress by someone 'optimising' the write path: the worker
        only ever constructs a new row -- there is no query for an existing one to mutate."""
        import inspect

        from app.domain import indexing

        source = inspect.getsource(indexing)
        assert "DocumentIngestion(" in source
        assert "select(DocumentIngestion" not in source
        assert ".get(DocumentIngestion" not in source


# =======================================================================================
# AC-04: a partially scanned PDF is distinguishable from a text document
# =======================================================================================
class TestLowTextDensity:
    """AC-04. The silent failure: a cover page with a text layer indexes SUCCESSFULLY."""

    def _upload_pdf(self, client, source_id: str, pages: list[str], title: str):
        return _upload(
            client,
            source_id,
            title=title,
            filename="policy.pdf",
            content=build_pdf_bytes(pages),
            mime="application/pdf",
        )

    def test_a_mostly_scanned_pdf_indexes_successfully_and_is_flagged(self, client):
        from app.domain.ingestion_lineage import WARNING_LOW_TEXT_DENSITY

        source = _create_source(client)
        response = self._upload_pdf(
            client, source["id"], ["Cover page only."] + [""] * 11, "Scanned Manual"
        )
        assert response.status_code == 201, response.text
        body = response.json()
        # This is the point of the criterion: the job SUCCEEDS. Without the warning there is
        # no signal at all that the policy text never made it into the index.
        assert body["index_job"]["status"] == "succeeded"

        (row,) = _lineage_rows(client, body["document"]["id"])
        assert WARNING_LOW_TEXT_DENSITY in row.warnings
        assert row.source_unit_kind == "page"
        assert row.source_unit_count == 12
        assert row.extracted_char_count == len("Cover page only.")

    def test_a_normal_text_pdf_is_not_flagged(self, client):
        from app.domain.ingestion_lineage import WARNING_LOW_TEXT_DENSITY

        source = _create_source(client)
        dense = "제15조 " + "연차 유급휴가는 근속 1년 이상 직원에게 부여한다. " * 12
        response = self._upload_pdf(client, source["id"], [dense.encode("ascii", "replace").decode()], "Text Manual")
        assert response.status_code == 201, response.text
        (row,) = _lineage_rows(client, response.json()["document"]["id"])
        assert WARNING_LOW_TEXT_DENSITY not in row.warnings
        assert row.source_unit_count == 1

    def test_the_raw_numbers_are_stored_so_the_threshold_can_be_disagreed_with(self, client):
        source = _create_source(client)
        response = self._upload_pdf(client, source["id"], ["Cover page only."] + [""] * 11, "Scan")
        (row,) = _lineage_rows(client, response.json()["document"]["id"])
        assert row.extracted_char_count is not None and row.source_unit_count is not None
        assert row.extracted_char_count / row.source_unit_count < 200

    def test_density_is_only_evaluated_for_page_counted_formats(self):
        """A short MD file is not a scanned document. Lines and paragraphs have no expected
        size, so applying a per-unit character threshold to them would be noise dressed as a
        finding."""
        from app.domain.ingestion_lineage import is_low_text_density
        from app.domain.parsers import UNIT_LINE, UNIT_PAGE, UNIT_PARAGRAPH, ExtractedDocument

        for unit_kind in (UNIT_LINE, UNIT_PARAGRAPH):
            assert not is_low_text_density(
                ExtractedDocument(text="hi", converter="c", source_unit_kind=unit_kind, source_unit_count=40)
            )
        assert is_low_text_density(
            ExtractedDocument(text="hi", converter="c", source_unit_kind=UNIT_PAGE, source_unit_count=40)
        )
        assert not is_low_text_density(
            ExtractedDocument(
                text="x" * 4000, converter="c", source_unit_kind=UNIT_PAGE, source_unit_count=4
            )
        )
        # No pages counted -> no ratio -> no claim (never a division by zero either).
        assert not is_low_text_density(
            ExtractedDocument(text="hi", converter="c", source_unit_kind=UNIT_PAGE, source_unit_count=0)
        )
        assert not is_low_text_density(None)


# =======================================================================================
# Failed attempts are recorded too -- an attempt that produced nothing is still a fact
# =======================================================================================
class TestFailedAttempts:
    def test_a_document_that_cannot_be_indexed_records_a_failed_row_with_no_counts(self, client):
        from app.domain.ingestion_lineage import EXTRACTION_FAILED

        source = _create_source(client)
        registered = client.post(
            "/api/v1/knowledge/documents",
            headers=_KM,
            json={
                "knowledge_source_id": source["id"],
                "title": "Deny-all",
                "object_uri": "object://deny.md",
                "checksum": "sha256-deny",
                "mime_type": "text/markdown",
                "access_groups": [],  # deny-all -> document_can_be_indexed is False
            },
        )
        assert registered.status_code == 201, registered.text
        document_id = registered.json()["id"]
        job = client.post(
            f"/api/v1/knowledge/documents/{document_id}/index-jobs",
            headers={"X-Agent-Forge-User": "admin", "X-Agent-Forge-Roles": "admin"},
            json={"source_text": POLICY_MD},
        )
        assert job.status_code == 201, job.text
        assert job.json()["error_code"] == "DOCUMENT_NOT_INDEXABLE"

        (row,) = _lineage_rows(client, document_id)
        assert row.extraction_status == EXTRACTION_FAILED
        assert row.source_mime_type == "text/markdown"
        # Nothing was converted, so nothing is claimed. NULL is "not observed", not zero.
        assert row.converted_mime_type is None
        assert row.chunk_count is None
        assert row.structured_chunk_count is None
        assert row.extracted_char_count is None
        assert row.warnings == []

    def test_a_successful_attempt_is_marked_ok(self, client):
        from app.domain.ingestion_lineage import EXTRACTION_OK

        source = _create_source(client)
        uploaded = _upload(
            client,
            source["id"],
            title="OK",
            filename="policy.md",
            content=POLICY_MD.encode("utf-8"),
            mime="text/markdown",
        )
        (row,) = _lineage_rows(client, uploaded.json()["document"]["id"])
        assert row.extraction_status == EXTRACTION_OK


# =======================================================================================
# The signal reaches the existing index-job read model
# =======================================================================================
class TestIndexJobReadModel:
    def test_the_index_job_response_carries_its_lineage(self, client):
        source = _create_source(client)
        uploaded = _upload(
            client,
            source["id"],
            title="취업규칙 (DOCX)",
            filename="policy.docx",
            content=build_docx_bytes(POLICY_MD),
            mime=DOCX_MIME,
        )
        job_id = uploaded.json()["index_job"]["id"]
        response = client.get(f"/api/v1/knowledge/index-jobs/{job_id}", headers=_KM)
        assert response.status_code == 200, response.text
        ingestion = response.json()["ingestion"]
        assert ingestion["converted_mime_type"] == "text/plain"
        assert ingestion["chunk_count"] == 1
        assert ingestion["structured_chunk_count"] == 0
        assert "NO_HEADINGS_DETECTED" in ingestion["warnings"]
        assert ingestion["index_job_id"] == job_id

    def test_a_queued_job_that_has_not_run_reports_null_rather_than_zero(self, client):
        """A job with no attempt yet must not look like an attempt that produced nothing."""
        source = _create_source(client)
        registered = client.post(
            "/api/v1/knowledge/documents",
            headers=_KM,
            json={
                "knowledge_source_id": source["id"],
                "title": "Queued",
                "object_uri": "object://queued.md",
                "checksum": "sha256-queued",
                "mime_type": "text/markdown",
                "access_groups": ["all-employees"],
            },
        )
        job = client.post(
            f"/api/v1/knowledge/documents/{registered.json()['id']}/index-jobs",
            headers=_KM,
            json={},  # no source_text -> queued, not run
        )
        assert job.status_code == 201, job.text
        assert job.json()["status"] == "queued"
        assert job.json()["ingestion"] is None


# =======================================================================================
# AC-05 / AC-06: the migration
# =======================================================================================
ALEMBIC_VERSION_NUM_MAX_LENGTH = 32


def _load_revision_module():
    """Import the revision file directly -- the pytest suite builds its schema with
    ``create_all`` and never runs Alembic, so without this the backfill would be covered by
    nothing until CI."""
    import importlib.util
    from pathlib import Path

    path = (
        Path(__file__).resolve().parents[1] / "alembic/versions/0007_document_ingestions.py"
    )
    spec = importlib.util.spec_from_file_location("revision_0007", path)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


class TestBackfillIsHonest:
    """AC-05."""

    def test_it_derives_the_two_counts_that_are_actually_derivable(self):
        module = _load_revision_module()
        from datetime import UTC, datetime

        rows = module._backfill_rows(
            document_mime_types={"doc-md": "text/markdown", "doc-docx": DOCX_MIME},
            chunk_rows=[
                # Postgres hands back a parsed list; SQLite hands back a JSON string. Both.
                ("doc-md", ["취업규칙", "제3장"]),
                ("doc-md", '["취업규칙", "제4장"]'),
                ("doc-md", "[]"),
                ("doc-docx", None),
                ("doc-docx", "[]"),
            ],
            created_at=datetime.now(UTC),
        )
        by_document = {row["document_id"]: row for row in rows}
        assert by_document["doc-md"]["chunk_count"] == 3
        assert by_document["doc-md"]["structured_chunk_count"] == 2
        assert by_document["doc-docx"]["chunk_count"] == 2
        assert by_document["doc-docx"]["structured_chunk_count"] == 0
        assert by_document["doc-docx"]["source_mime_type"] == DOCX_MIME

    def test_everything_else_is_marked_unknown_or_left_null(self):
        module = _load_revision_module()
        from datetime import UTC, datetime

        (row,) = module._backfill_rows(
            document_mime_types={"doc-1": "application/pdf"},
            chunk_rows=[("doc-1", "[]")],
            created_at=datetime.now(UTC),
        )
        # THE marker: no existing document is recorded as having a verified extraction.
        assert row["extraction_status"] == "unknown"
        assert row["index_job_id"] is None
        assert row["converted_mime_type"] is None
        assert row["converter_chain"] is None
        assert row["extracted_char_count"] is None
        assert row["source_unit_kind"] is None
        assert row["source_unit_count"] is None
        assert row["warnings"] == []

    def test_a_document_with_no_chunks_gets_no_row(self):
        """No evidence an attempt ever happened -> no record claiming one did."""
        module = _load_revision_module()
        from datetime import UTC, datetime

        assert (
            module._backfill_rows(
                document_mime_types={"doc-1": "text/markdown"},
                chunk_rows=[],
                created_at=datetime.now(UTC),
            )
            == []
        )

    def test_unparseable_section_paths_count_as_absent_not_present(self):
        """The cautious direction: under-claim structure, never over-claim it."""
        module = _load_revision_module()
        assert module._section_path_is_present('["a"]') is True
        assert module._section_path_is_present(["a"]) is True
        assert module._section_path_is_present("[]") is False
        assert module._section_path_is_present(None) is False
        assert module._section_path_is_present("not json at all") is False
        assert module._section_path_is_present(b'["a"]') is True

    def test_the_revision_cannot_write_a_verified_status(self):
        import re
        from pathlib import Path

        text = (
            Path(__file__).resolve().parents[1] / "alembic/versions/0007_document_ingestions.py"
        ).read_text(encoding="utf-8")
        assert 'UNKNOWN_EXTRACTION_STATUS = "unknown"' in text
        assert '"extraction_status": UNKNOWN_EXTRACTION_STATUS' in text
        # The prohibited action, asserted structurally: the revision never assigns a STRING
        # LITERAL to extraction_status, so the only value it can write is the named unknown.
        assert re.search(r'"extraction_status":\s*["\']', text) is None


class TestMigrationParity:
    """AC-06 plus model/migration agreement (nothing else in the suite runs Alembic)."""

    from pathlib import Path as _Path

    REVISION = _Path(__file__).resolve().parents[1] / "alembic/versions/0007_document_ingestions.py"

    def test_it_chains_to_0006_and_fits_alembics_version_column(self):
        text = self.REVISION.read_text(encoding="utf-8")
        assert 'revision: str = "0007_document_ingestions"' in text
        assert 'down_revision: str | None = "0006_source_class_defaults"' in text
        assert len("0007_document_ingestions") <= ALEMBIC_VERSION_NUM_MAX_LENGTH

    def test_every_model_column_is_created_by_the_revision(self):
        from app.domain.models import DocumentIngestion

        text = self.REVISION.read_text(encoding="utf-8")
        upgrade_body, _, downgrade_body = text.partition("def downgrade()")
        for column in DocumentIngestion.__table__.columns.keys():
            assert f'"{column}"' in upgrade_body, f"0007 does not create {column}"
        assert 'op.drop_table(TABLE_NAME)' in downgrade_body

    def test_the_table_is_append_only_by_construction(self):
        """Nothing UNIQUE over document_id or index_job_id: the schema itself must permit many
        attempts per document, or AC-02 would hold by luck rather than by design."""
        from sqlalchemy import PrimaryKeyConstraint, UniqueConstraint

        from app.domain.models import DocumentIngestion

        table = DocumentIngestion.__table__
        for constraint in table.constraints:
            if not isinstance(constraint, UniqueConstraint | PrimaryKeyConstraint):
                continue
            columns = {column.name for column in constraint.columns}
            assert columns == {"id"}, f"unexpected uniqueness over {sorted(columns)}"
        for index in table.indexes:
            assert not index.unique, f"unique index {index.name} would cap attempts per document"

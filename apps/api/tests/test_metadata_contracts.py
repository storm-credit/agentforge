import importlib.util
from collections.abc import Iterator

import pytest


RUNTIME_DEPS = ("fastapi", "httpx", "pydantic_settings", "sqlalchemy")


def runtime_deps_available() -> bool:
    return all(importlib.util.find_spec(package) for package in RUNTIME_DEPS)


pytestmark = pytest.mark.skipif(
    not runtime_deps_available(),
    reason="Runtime dependencies are not installed",
)


@pytest.fixture
def client():
    from fastapi.testclient import TestClient
    from sqlalchemy import create_engine
    from sqlalchemy.orm import Session, sessionmaker
    from sqlalchemy.pool import StaticPool

    from app.core.database import Base, get_db
    from app.domain import models  # noqa: F401
    from app.main import create_app

    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session = sessionmaker(bind=engine, autoflush=False, expire_on_commit=False)
    Base.metadata.create_all(bind=engine)

    def override_get_db() -> Iterator[Session]:
        db = testing_session()
        try:
            yield db
        finally:
            db.close()

    app = create_app()
    app.dependency_overrides[get_db] = override_get_db

    with TestClient(app) as test_client:
        yield test_client

    Base.metadata.drop_all(bind=engine)


def test_agent_and_version_contract(client):
    agent_response = client.post(
        "/api/v1/agents",
        headers={
            "X-Agent-Forge-User": "agent-owner",
            "X-Agent-Forge-Department": "Operations",
        },
        json={
            "name": "Policy Assistant",
            "purpose": "Answer internal policy questions with citations.",
            "owner_department": "Operations",
        },
    )

    assert agent_response.status_code == 201
    agent = agent_response.json()
    assert agent["id"]
    assert agent["name"] == "Policy Assistant"
    assert agent["status"] == "draft"

    detail_response = client.get(f"/api/v1/agents/{agent['id']}")

    assert detail_response.status_code == 200
    assert detail_response.json()["id"] == agent["id"]

    update_response = client.patch(
        f"/api/v1/agents/{agent['id']}",
        headers={"X-Agent-Forge-User": "agent-owner"},
        json={"purpose": "Answer internal policy questions with traceable citations."},
    )

    assert update_response.status_code == 200
    assert update_response.json()["purpose"] == (
        "Answer internal policy questions with traceable citations."
    )

    version_response = client.post(
        "/api/v1/agents/versions",
        headers={"X-Agent-Forge-User": "agent-owner"},
        json={
            "agent_id": agent["id"],
            "version": 1,
            "config": {
                "citation_required": True,
                "knowledge_source_ids": [],
            },
        },
    )

    assert version_response.status_code == 201
    version = version_response.json()
    assert version["agent_id"] == agent["id"]
    assert version["version"] == 1
    assert version["created_by"] == "agent-owner"

    versions_response = client.get(f"/api/v1/agents/{agent['id']}/versions")

    assert versions_response.status_code == 200
    assert [item["id"] for item in versions_response.json()] == [version["id"]]

    validate_response = client.post(
        f"/api/v1/agents/versions/{version['id']}/validate",
        headers={"X-Agent-Forge-User": "agent-owner", "X-Agent-Forge-Roles": "admin"},
        json={"reason": "Contract test validation"},
    )

    assert validate_response.status_code == 200
    assert validate_response.json()["status"] == "validated"

    publish_response = client.post(
        f"/api/v1/agents/versions/{version['id']}/publish",
        headers={"X-Agent-Forge-User": "platform-admin", "X-Agent-Forge-Roles": "admin"},
        json={"reason": "Contract test publish"},
    )

    assert publish_response.status_code == 200
    assert publish_response.json()["status"] == "published"

    list_response = client.get("/api/v1/agents")

    assert list_response.status_code == 200
    listed_agents = list_response.json()
    assert [item["id"] for item in listed_agents] == [agent["id"]]
    assert listed_agents[0]["status"] == "published"


_ADMIN = {"X-Agent-Forge-User": "ops", "X-Agent-Forge-Roles": "admin"}
_DEVELOPER = {"X-Agent-Forge-User": "dev", "X-Agent-Forge-Roles": "developer"}


def test_audit_events_query_admin_only_and_self_audited(client):
    # non-admin is denied
    assert client.get("/api/v1/audit/events", headers=_DEVELOPER).status_code == 403

    # generate an auditable event
    src = client.post(
        "/api/v1/knowledge/sources",
        json={"name": "Audit Src", "description": "x", "owner_department": "Operations"},
    ).json()

    resp = client.get(
        "/api/v1/audit/events",
        headers=_ADMIN,
        params={"event_type": "knowledge_source.created"},
    )
    assert resp.status_code == 200
    events = resp.json()
    assert any(e["target_id"] == src["id"] for e in events)
    assert all(e["event_type"] == "knowledge_source.created" for e in events)

    # the query itself is audited (audit_log.viewed)
    viewed = client.get(
        "/api/v1/audit/events", headers=_ADMIN, params={"event_type": "audit_log.viewed"}
    )
    assert viewed.status_code == 200
    assert len(viewed.json()) >= 1


def test_archive_document_excludes_from_retrieval_and_list(client):
    document = _create_indexable_document(client)
    client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        json={"source_text": "# Remote Work\n\nremote work after manager approval."},
    )

    def _doc_ids():
        return [d["id"] for d in client.get("/api/v1/knowledge/documents").json()]

    def _preview_titles():
        resp = client.post(
            "/api/v1/knowledge/retrieval/preview",
            json={"query": "remote work approval", "top_k": 10},
        )
        return [h["title"] for h in resp.json()["hits"]]

    assert document["id"] in _doc_ids()
    assert document["title"] in _preview_titles()

    # non-admin cannot archive
    assert client.delete(
        f"/api/v1/knowledge/documents/{document['id']}", headers=_DEVELOPER
    ).status_code == 403

    # admin archives -> status archived
    resp = client.delete(
        f"/api/v1/knowledge/documents/{document['id']}?reason=stale doc",
        headers=_ADMIN,
    )
    assert resp.status_code == 200
    assert resp.json()["status"] == "archived"

    # excluded from list + retrieval
    assert document["id"] not in _doc_ids()
    assert document["title"] not in _preview_titles()

    # audited
    events = client.get(
        "/api/v1/audit/events?event_type=document.archived", headers=_ADMIN
    ).json()
    assert any(e["target_id"] == document["id"] for e in events)


def test_document_list_and_chunks_scoped_by_acl(client):
    source = client.post(
        "/api/v1/knowledge/sources",
        json={"name": "Scoped Src", "description": "x", "owner_department": "Security"},
    ).json()
    doc = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "HR Restricted Doc",
            "object_uri": "object://hr-restricted.md",
            "checksum": "sha256-hr-restricted",
            "mime_type": "text/markdown",
            "confidentiality_level": "restricted",
            "access_groups": ["department:HR"],
        },
    ).json()
    client.post(
        f"/api/v1/knowledge/documents/{doc['id']}/index-jobs",
        json={"source_text": "# HR\n\nrestricted leave exception details."},
    )

    finance = {"X-Agent-Forge-Department": "Finance", "X-Agent-Forge-Clearance": "internal"}
    hr = {
        "X-Agent-Forge-Department": "HR",
        "X-Agent-Forge-Groups": "department:HR",
        "X-Agent-Forge-Clearance": "restricted",
    }

    # chunks: Finance (no access) blocked; HR allowed; admin allowed
    assert client.get(
        f"/api/v1/knowledge/documents/{doc['id']}/chunks", headers=finance
    ).status_code == 403
    assert client.get(
        f"/api/v1/knowledge/documents/{doc['id']}/chunks", headers=hr
    ).status_code == 200
    assert client.get(
        f"/api/v1/knowledge/documents/{doc['id']}/chunks", headers=_ADMIN
    ).status_code == 200

    # list: Finance does not see the restricted doc; admin does
    finance_ids = [d["id"] for d in client.get("/api/v1/knowledge/documents", headers=finance).json()]
    admin_ids = [d["id"] for d in client.get("/api/v1/knowledge/documents", headers=_ADMIN).json()]
    assert doc["id"] not in finance_ids
    assert doc["id"] in admin_ids


def test_audit_events_pagination(client):
    for i in range(3):
        client.post(
            "/api/v1/knowledge/sources",
            json={"name": f"Pg {i}", "description": "x", "owner_department": "Operations"},
        )
    page = client.get("/api/v1/audit/events", headers=_ADMIN, params={"limit": 2})
    assert page.status_code == 200
    assert len(page.json()) == 2


def test_privileged_mutations_require_admin_role(client):
    # set up an agent + draft version + an indexed document as admin
    agent = client.post(
        "/api/v1/agents",
        json={"name": "RBAC Agent", "purpose": "rbac.", "owner_department": "Operations"},
    ).json()
    version = client.post(
        "/api/v1/agents/versions", json={"agent_id": agent["id"], "config": {}}
    ).json()
    source = client.post(
        "/api/v1/knowledge/sources",
        json={"name": "RBAC Src", "description": "x", "owner_department": "Operations"},
    ).json()
    doc = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "RBAC Doc",
            "object_uri": "object://rbac.md",
            "checksum": "sha256-rbac",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": ["all-employees"],
        },
    ).json()

    # developer role is rejected on each privileged mutation
    assert client.post(
        f"/api/v1/agents/versions/{version['id']}/validate",
        headers=_DEVELOPER,
        json={"reason": "x"},
    ).status_code == 403
    assert client.post(
        f"/api/v1/agents/versions/{version['id']}/publish",
        headers=_DEVELOPER,
        json={"reason": "x"},
    ).status_code == 403
    assert client.patch(
        f"/api/v1/knowledge/documents/{doc['id']}/acl",
        headers=_DEVELOPER,
        json={"access_groups": ["public"], "confidentiality_level": "public", "reason": "x"},
    ).status_code == 403

    # admin role is allowed
    assert client.post(
        f"/api/v1/agents/versions/{version['id']}/validate",
        headers=_ADMIN,
        json={"reason": "ok"},
    ).status_code == 200
    assert client.post(
        f"/api/v1/agents/versions/{version['id']}/publish",
        headers=_ADMIN,
        json={"reason": "ok"},
    ).status_code == 200
    assert client.patch(
        f"/api/v1/knowledge/documents/{doc['id']}/acl",
        headers=_ADMIN,
        json={"access_groups": ["department:HR"], "confidentiality_level": "internal", "reason": "ok"},
    ).status_code == 200


def test_agent_versions_autonumber_on_create(client):
    agent = client.post(
        "/api/v1/agents",
        json={
            "name": "Versioned Assistant",
            "purpose": "Auto-number versions.",
            "owner_department": "Operations",
        },
    ).json()

    first = client.post(
        "/api/v1/agents/versions",
        json={"agent_id": agent["id"], "config": {"citation_required": True}},
    )
    second = client.post(
        "/api/v1/agents/versions",
        json={"agent_id": agent["id"], "config": {"citation_required": True}},
    )

    assert first.status_code == 201
    assert second.status_code == 201
    assert first.json()["version"] == 1
    assert second.json()["version"] == 2


def test_rollback_republishes_older_version(client):
    agent = client.post(
        "/api/v1/agents",
        json={"name": "Rollback Agent", "purpose": "rollback.", "owner_department": "Operations"},
    ).json()
    v1 = client.post(
        "/api/v1/agents/versions", json={"agent_id": agent["id"], "config": {}}
    ).json()
    v2 = client.post(
        "/api/v1/agents/versions", json={"agent_id": agent["id"], "config": {}}
    ).json()

    client.post(f"/api/v1/agents/versions/{v1['id']}/publish", headers=_ADMIN, json={"reason": "publish v1"})
    client.post(f"/api/v1/agents/versions/{v2['id']}/publish", headers=_ADMIN, json={"reason": "publish v2"})

    # rollback = re-publish the older version; it supersedes the current one
    rollback = client.post(
        f"/api/v1/agents/versions/{v1['id']}/publish", headers=_ADMIN, json={"reason": "rollback to v1"}
    )
    assert rollback.status_code == 200
    assert rollback.json()["status"] == "published"

    versions = client.get(f"/api/v1/agents/{agent['id']}/versions").json()
    status_by_version = {v["version"]: v["status"] for v in versions}
    assert status_by_version[1] == "published"
    assert status_by_version[2] == "superseded"


def test_knowledge_source_and_document_contract(client):
    source_response = client.post(
        "/api/v1/knowledge/sources",
        headers={
            "X-Agent-Forge-User": "knowledge-manager",
            "X-Agent-Forge-Department": "Operations",
        },
        json={
            "name": "Pilot Policies",
            "description": "Synthetic policy documents for Sprint 0.",
            "owner_department": "Operations",
            "default_confidentiality_level": "internal",
        },
    )

    assert source_response.status_code == 201
    source = source_response.json()
    assert source["id"]
    assert source["default_confidentiality_level"] == "internal"

    document_response = client.post(
        "/api/v1/knowledge/documents",
        headers={"X-Agent-Forge-User": "knowledge-manager"},
        json={
            "knowledge_source_id": source["id"],
            "title": "Remote Work Policy",
            "object_uri": "object://synthetic/hr/remote-work.md",
            "checksum": "sha256-synthetic-remote-work",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": ["all-employees"],
            "status": "registered",
            "effective_date": "2026-05-09",
        },
    )

    assert document_response.status_code == 201
    document = document_response.json()
    assert document["knowledge_source_id"] == source["id"]
    assert document["access_groups"] == ["all-employees"]

    list_response = client.get("/api/v1/knowledge/documents")

    assert list_response.status_code == 200
    assert [item["id"] for item in list_response.json()] == [document["id"]]


def test_retrieval_preview_applies_document_acl(client):
    source_response = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Sprint 1 ACL Corpus",
            "description": "Synthetic ACL corpus.",
            "owner_department": "Security",
        },
    )
    source = source_response.json()

    public_response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "Company Holiday Policy",
            "object_uri": "object://synthetic/company/holiday.md",
            "checksum": "sha256-holiday",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": ["all-employees"],
        },
    )
    restricted_response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "HR Leave Exception Policy",
            "object_uri": "object://synthetic/hr/leave-exception.md",
            "checksum": "sha256-leave-exception",
            "mime_type": "text/markdown",
            "confidentiality_level": "restricted",
            "access_groups": ["department:HR"],
        },
    )
    no_acl_response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "Unclassified Draft Policy",
            "object_uri": "object://synthetic/drafts/unclassified.md",
            "checksum": "sha256-unclassified",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": [],
        },
    )

    assert public_response.status_code == 201
    assert restricted_response.status_code == 201
    assert no_acl_response.status_code == 201

    finance_response = client.post(
        "/api/v1/knowledge/retrieval/preview",
        headers={
            "X-Agent-Forge-Department": "Finance",
            "X-Agent-Forge-Clearance": "internal",
        },
        json={
            "query": "leave policy",
            "knowledge_source_ids": [source["id"]],
            "top_k": 10,
        },
    )

    assert finance_response.status_code == 200
    finance_hits = finance_response.json()["hits"]
    assert [hit["title"] for hit in finance_hits] == ["Company Holiday Policy"]
    assert finance_response.json()["denied_count"] == 2

    hr_response = client.post(
        "/api/v1/knowledge/retrieval/preview",
        headers={
            "X-Agent-Forge-Department": "HR",
            "X-Agent-Forge-Clearance": "restricted",
        },
        json={
            "query": "leave policy",
            "knowledge_source_ids": [source["id"]],
            "top_k": 10,
        },
    )

    assert hr_response.status_code == 200
    hr_titles = [hit["title"] for hit in hr_response.json()["hits"]]
    assert hr_titles == ["HR Leave Exception Policy", "Company Holiday Policy"]
    assert "Unclassified Draft Policy" not in hr_titles


def test_index_job_creates_txt_md_chunks_and_preview_uses_chunk_citations(client):
    source_response = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Sprint 1 Parser Corpus",
            "description": "Synthetic parser corpus.",
            "owner_department": "Operations",
        },
    )
    source = source_response.json()
    document_response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "Remote Work Policy",
            "object_uri": "object://synthetic/ops/remote-work.md",
            "checksum": "sha256-remote-work-parser",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": ["all-employees"],
            "effective_date": "2026-05-10",
        },
    )
    document = document_response.json()

    job_response = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        headers={"X-Agent-Forge-User": "indexer"},
        json={
            "source_text": (
                "# Remote Work\n\n"
                "Company-wide remote work rules.\n\n"
                "## Eligibility\n\n"
                "Employees may request remote work after manager approval."
            ),
            "chunking": {"strategy": "line-heading", "chunk_size": 900, "chunk_overlap": 0},
        },
    )

    assert job_response.status_code == 201
    job = job_response.json()
    assert job["status"] == "succeeded"
    assert job["stage"] == "upsert"
    assert job["chunk_count"] == 2
    assert job["created_by"] == "indexer"

    detail_response = client.get(f"/api/v1/knowledge/index-jobs/{job['id']}")

    assert detail_response.status_code == 200
    assert detail_response.json()["id"] == job["id"]

    chunks_response = client.get(f"/api/v1/knowledge/documents/{document['id']}/chunks")

    assert chunks_response.status_code == 200
    chunks = chunks_response.json()
    assert len(chunks) == 2
    assert "content" not in chunks[0]
    assert chunks[1]["section_path"] == ["Remote Work", "Eligibility"]
    assert chunks[1]["citation_locator"] == (
        "Remote Work Policy / Remote Work > Eligibility / lines 7-7"
    )
    assert chunks[1]["vector_ref"].startswith("fake-vector:none-smoke:")
    assert chunks[1]["acl_snapshot"]["access_groups"] == ["all-employees"]

    preview_response = client.post(
        "/api/v1/knowledge/retrieval/preview",
        json={
            "query": "manager approval",
            "knowledge_source_ids": [source["id"]],
            "top_k": 1,
        },
    )

    assert preview_response.status_code == 200
    hit = preview_response.json()["hits"][0]
    assert hit["chunk_id"] == chunks[1]["id"]
    assert hit["citation_locator"] == chunks[1]["citation_locator"]


def test_pdf_upload_extracts_text_and_indexes_chunks(client):
    source = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Binary Upload Corpus",
            "description": "PDF and DOCX upload corpus.",
            "owner_department": "Operations",
        },
    ).json()

    response = client.post(
        "/api/v1/knowledge/documents/upload",
        data={
            "knowledge_source_id": source["id"],
            "title": "Remote Work PDF",
            "confidentiality_level": "internal",
            "access_groups": "all-employees",
            "effective_date": "2026-06-12",
        },
        files={
            "file": (
                "remote-work.pdf",
                _minimal_pdf_bytes(["Remote work requires manager approval."]),
                "application/pdf",
            )
        },
        headers={"X-Agent-Forge-User": "binary-uploader"},
    )

    assert response.status_code == 201
    payload = response.json()
    document = payload["document"]
    job = payload["index_job"]
    assert document["mime_type"] == "application/pdf"
    assert document["status"] == "indexed"
    assert document["checksum"].startswith("sha256-")
    assert document["object_uri"] == "upload://remote-work.pdf"
    assert job["status"] == "succeeded"
    assert job["chunk_count"] == 1
    assert job["config"]["source"] == "uploaded_file"
    assert job["created_by"] == "binary-uploader"

    chunks = client.get(f"/api/v1/knowledge/documents/{document['id']}/chunks").json()
    assert len(chunks) == 1
    assert chunks[0]["acl_snapshot"]["access_groups"] == ["all-employees"]
    assert "lines" in chunks[0]["citation_locator"]


def test_docx_upload_extracts_text_and_indexes_chunks(client):
    source = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "DOCX Upload Corpus",
            "description": "DOCX upload corpus.",
            "owner_department": "Operations",
        },
    ).json()

    response = client.post(
        "/api/v1/knowledge/documents/upload",
        data={
            "knowledge_source_id": source["id"],
            "title": "Travel Policy DOCX",
            "confidentiality_level": "internal",
            "access_groups": "all-employees",
        },
        files={
            "file": (
                "travel-policy.docx",
                _docx_bytes("Travel Policy", "Travel stipend is fifty dollars per day."),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 201
    payload = response.json()
    document = payload["document"]
    job = payload["index_job"]
    assert document["mime_type"] == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert document["status"] == "indexed"
    assert job["status"] == "succeeded"
    assert job["chunk_count"] >= 1

    chunks = client.get(f"/api/v1/knowledge/documents/{document['id']}/chunks").json()
    assert len(chunks) >= 1
    assert chunks[0]["citation_locator"].startswith("Travel Policy DOCX /")


def test_binary_upload_rejects_unsupported_mime_type(client):
    source = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Unsupported Upload Corpus",
            "description": "Unsupported upload corpus.",
            "owner_department": "Operations",
        },
    ).json()

    response = client.post(
        "/api/v1/knowledge/documents/upload",
        data={
            "knowledge_source_id": source["id"],
            "title": "Unsupported",
            "confidentiality_level": "internal",
            "access_groups": "all-employees",
        },
        files={"file": ("unknown.bin", b"unknown", "application/octet-stream")},
    )

    assert response.status_code == 415
    assert response.json()["detail"] == "Unsupported file type"


def test_index_job_fails_closed_for_missing_acl(client):
    source_response = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Missing ACL Corpus",
            "description": "Documents without ACL must not become searchable.",
            "owner_department": "Security",
        },
    )
    source = source_response.json()
    document_response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "Unclassified Draft Policy",
            "object_uri": "object://synthetic/security/no-acl.md",
            "checksum": "sha256-no-acl-parser",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": [],
        },
    )
    document = document_response.json()

    job_response = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        json={"source_text": "# Draft\n\nThis must not become searchable."},
    )

    assert job_response.status_code == 201
    job = job_response.json()
    assert job["status"] == "failed"
    assert job["error_code"] == "DOCUMENT_NOT_INDEXABLE"
    assert job["chunk_count"] == 0

    chunks_response = client.get(
        f"/api/v1/knowledge/documents/{document['id']}/chunks", headers=_ADMIN
    )

    assert chunks_response.status_code == 200
    assert chunks_response.json() == []


def test_index_job_rejects_unknown_document(client):
    response = client.post(
        "/api/v1/knowledge/documents/missing-document/index-jobs",
        json={"source_text": "Missing document."},
    )

    assert response.status_code == 404
    assert response.json()["detail"] == "Document not found"


def _create_indexable_document(client) -> dict:
    source = client.post(
        "/api/v1/knowledge/sources",
        json={
            "name": "Queued Worker Corpus",
            "description": "Documents indexed through the queued worker stub.",
            "owner_department": "Operations",
        },
    ).json()
    return client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": source["id"],
            "title": "Async Remote Work Policy",
            "object_uri": "object://synthetic/ops/async-remote-work.md",
            "checksum": "sha256-async-remote-work",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
            "access_groups": ["all-employees"],
            "effective_date": "2026-05-10",
        },
    ).json()


def test_queued_index_job_is_processed_by_worker(client):
    document = _create_indexable_document(client)

    queued_response = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        json={"chunking": {"strategy": "line-heading", "chunk_size": 900, "chunk_overlap": 0}},
    )
    assert queued_response.status_code == 201
    queued_job = queued_response.json()
    assert queued_job["status"] == "queued"
    assert queued_job["config"]["source"] == "object_store"

    process_response = client.post(
        f"/api/v1/knowledge/index-jobs/{queued_job['id']}/process",
        json={
            "source_text": (
                "# Remote Work\n\n"
                "Company-wide remote work rules.\n\n"
                "## Eligibility\n\n"
                "Employees may request remote work after manager approval."
            )
        },
    )
    assert process_response.status_code == 200
    processed_job = process_response.json()
    assert processed_job["id"] == queued_job["id"]
    assert processed_job["status"] == "succeeded"
    assert processed_job["stage"] == "upsert"
    assert processed_job["chunk_count"] == 2

    chunks = client.get(f"/api/v1/knowledge/documents/{document['id']}/chunks").json()
    assert len(chunks) == 2


def test_queued_index_job_without_content_fails_closed(client):
    document = _create_indexable_document(client)

    queued_job = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        json={},
    ).json()
    assert queued_job["status"] == "queued"

    process_response = client.post(
        f"/api/v1/knowledge/index-jobs/{queued_job['id']}/process",
        json={},
    )
    assert process_response.status_code == 200
    processed_job = process_response.json()
    assert processed_job["status"] == "failed"
    assert processed_job["error_code"] == "SOURCE_CONTENT_UNAVAILABLE"
    assert processed_job["chunk_count"] == 0

    chunks = client.get(
        f"/api/v1/knowledge/documents/{document['id']}/chunks", headers=_ADMIN
    ).json()
    assert chunks == []


def test_queued_index_fetches_content_from_object_store(client, monkeypatch):
    from app.core.config import get_settings
    from app.infra.object_store import document_object_key, get_object_store

    monkeypatch.setenv("AGENT_FORGE_OBJECT_STORE_BACKEND", "memory")
    get_settings.cache_clear()
    get_object_store.cache_clear()
    try:
        document = _create_indexable_document(client)
        store = get_object_store()
        store.put(
            document_object_key(document["id"]),
            "# Stored Policy\n\nThis content was fetched from the object store.".encode("utf-8"),
        )

        queued = client.post(
            f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
            json={},
        ).json()
        assert queued["status"] == "queued"

        # No inline source_text: the worker must fetch the original from object storage
        # instead of failing closed with SOURCE_CONTENT_UNAVAILABLE.
        processed = client.post(
            f"/api/v1/knowledge/index-jobs/{queued['id']}/process",
            json={},
        )
        assert processed.status_code == 200
        body = processed.json()
        assert body["status"] == "succeeded"
        assert body["chunk_count"] >= 1
    finally:
        get_settings.cache_clear()
        get_object_store.cache_clear()


def test_queued_index_without_object_store_still_fails_closed(client):
    # Default backend = none -> no object store -> queued job without content fails closed.
    document = _create_indexable_document(client)
    queued = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs", json={}
    ).json()
    processed = client.post(
        f"/api/v1/knowledge/index-jobs/{queued['id']}/process", json={}
    ).json()
    assert processed["status"] == "failed"
    assert processed["error_code"] == "SOURCE_CONTENT_UNAVAILABLE"


def test_process_rejects_non_queued_job(client):
    document = _create_indexable_document(client)

    job = client.post(
        f"/api/v1/knowledge/documents/{document['id']}/index-jobs",
        json={"source_text": "# Title\n\nBody."},
    ).json()
    assert job["status"] == "succeeded"

    response = client.post(
        f"/api/v1/knowledge/index-jobs/{job['id']}/process",
        json={"source_text": "# Title\n\nBody."},
    )
    assert response.status_code == 409
    assert response.json()["detail"] == "Index job is not queued"


def test_document_registration_rejects_unknown_source(client):
    response = client.post(
        "/api/v1/knowledge/documents",
        json={
            "knowledge_source_id": "missing-source",
            "title": "Unknown Source Document",
            "object_uri": "object://synthetic/missing.md",
            "checksum": "sha256-missing",
            "mime_type": "text/markdown",
            "confidentiality_level": "internal",
        },
    )

    assert response.status_code == 404
    assert response.json()["detail"] == "Knowledge source not found"


def _docx_bytes(heading: str, paragraph: str) -> bytes:
    from io import BytesIO

    from docx import Document

    document = Document()
    document.add_heading(heading, level=1)
    document.add_paragraph(paragraph)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _minimal_pdf_bytes(lines: list[str]) -> bytes:
    text_ops = ["BT", "/F1 18 Tf", "72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            text_ops.append("0 -24 Td")
        safe_line = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        text_ops.append(f"({safe_line}) Tj")
    text_ops.append("ET")
    stream = "\n".join(text_ops).encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n"
        + stream
        + b"\nendstream",
    ]
    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_index, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_index} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")
    xref_at = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_at}\n%%EOF\n".encode("ascii")
    )
    return bytes(output)

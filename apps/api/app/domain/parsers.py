import hashlib
import re
from dataclasses import dataclass
from importlib.metadata import PackageNotFoundError, version as _package_version
from io import BytesIO


PARSER_VERSION = "txt-md-parser/0.1.0"
CHUNKER_VERSION = "token-window-overlap-chunker/0.2.0"
SUPPORTED_TEXT_MIME_TYPES = {"text/plain", "text/markdown", "text/x-markdown"}
PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
SUPPORTED_BINARY_MIME_TYPES = {PDF_MIME_TYPE, DOCX_MIME_TYPE}
SUPPORTED_DOCUMENT_MIME_TYPES = SUPPORTED_TEXT_MIME_TYPES | SUPPORTED_BINARY_MIME_TYPES
MAX_EXTRACT_BYTES = 10 * 1024 * 1024
MAX_PDF_PAGES = 50


#: What produced the text when no format conversion happened at all (TXT/MD upload bytes).
TEXT_IDENTITY_CONVERTER = "identity/decode-utf8"

#: The unit an extractor counts, so ``extracted_char_count / source_unit_count`` means
#: something. Only ``page`` supports the low-density signal: a page is a fixed-size physical
#: surface, so "too few characters per page" is evidence of an image-only (scanned) page. A
#: paragraph or a line has no expected size, so the same ratio would be noise.
UNIT_PAGE = "page"
UNIT_PARAGRAPH = "paragraph"
UNIT_LINE = "line"


class DocumentExtractionError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code

    @property
    def error_code(self) -> str:
        return self.code


@dataclass(frozen=True)
class ParsedChunk:
    chunk_id: str
    chunk_index: int
    content: str
    content_hash: str
    chunk_hash: str
    token_count: int
    line_start: int
    line_end: int
    section_path: tuple[str, ...]
    citation_locator: str
    parser_version: str
    chunker_version: str


def parse_txt_md_document(
    *,
    document_id: str,
    document_version: str,
    title: str,
    mime_type: str,
    source_text: str,
    chunk_size: int = 900,  # deprecated: retained for call-site compatibility, unused
    target_tokens: int = 320,
    overlap_tokens: int = 50,
) -> list[ParsedChunk]:
    """Chunk text with a heading-bounded sliding token (eojeol) window.

    Words within one heading section accumulate up to ``target_tokens``; adjacent
    chunks overlap by ``overlap_tokens`` so evidence is not lost at boundaries. The
    window never crosses a heading boundary, preserving ``section_path`` and the
    line-range citation locator. Token counts use the whitespace-word proxy
    (``len(text.split())``) — see the chunk-overlap design doc for the caveat.
    """
    if mime_type not in SUPPORTED_TEXT_MIME_TYPES:
        raise ValueError(f"Unsupported text parser MIME type: {mime_type}")

    normalized_text = source_text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized_text.split("\n")
    is_markdown = mime_type in {"text/markdown", "text/x-markdown"}
    chunks: list[ParsedChunk] = []
    heading_path: list[str] = []
    section_words: list[tuple[str, int]] = []
    section_heading_path: tuple[str, ...] = ()

    def flush_section() -> None:
        nonlocal section_words
        _emit_window_chunks(
            chunks=chunks,
            document_id=document_id,
            document_version=document_version,
            title=title,
            section_path=section_heading_path,
            words=section_words,
            target_tokens=target_tokens,
            overlap_tokens=overlap_tokens,
        )
        section_words = []

    for line_number, raw_line in enumerate(lines, start=1):
        heading = _parse_markdown_heading(raw_line) if is_markdown else None
        if heading is not None:
            flush_section()
            level, text = heading
            heading_path = heading_path[: level - 1] + [text]
            continue

        if not raw_line.strip():
            continue

        if not section_words:
            section_heading_path = tuple(heading_path)
        section_words.extend((word, line_number) for word in raw_line.split())

    flush_section()
    return chunks


def _emit_window_chunks(
    *,
    chunks: list[ParsedChunk],
    document_id: str,
    document_version: str,
    title: str,
    section_path: tuple[str, ...],
    words: list[tuple[str, int]],
    target_tokens: int,
    overlap_tokens: int,
) -> None:
    n = len(words)
    if n == 0:
        return
    target = max(1, target_tokens)
    step = max(1, target - max(0, overlap_tokens))
    i = 0
    while i < n:
        window = words[i : i + target]
        chunks.append(
            _build_chunk(
                document_id=document_id,
                document_version=document_version,
                title=title,
                chunk_index=len(chunks),
                content=" ".join(word for word, _ in window),
                line_start=window[0][1],
                line_end=window[-1][1],
                section_path=section_path,
            )
        )
        if i + target >= n:
            break
        i += step


@dataclass(frozen=True)
class ExtractedDocument:
    """What one extraction produced, alongside the text itself.

    The extra fields exist so an index attempt can RECORD what the conversion produced
    (WO-2026-08-14-INGESTION-INSTRUMENTATION). They are observations of the extractor's own
    inputs -- which library ran, how many pages/paragraphs/lines it saw, how many characters
    came out -- and nothing here is a judgement about quality.

    ``text`` is the same string ``extract_text_from_bytes`` has always returned: that function
    is now a wrapper around this one rather than a parallel implementation, so the two cannot
    drift apart.
    """

    text: str
    #: Extractor and its installed version, e.g. ``pypdf/6.14.2``.
    converter: str
    #: One of ``UNIT_PAGE`` / ``UNIT_PARAGRAPH`` / ``UNIT_LINE``.
    source_unit_kind: str
    #: How many of those units the extractor saw -- INCLUDING ones that yielded no text.
    #: For PDF this is the page count from the document itself, which is exactly what makes
    #: a partially scanned file detectable: the pages are there, the characters are not.
    source_unit_count: int


def extract_document(
    *,
    mime_type: str,
    content: bytes,
    max_bytes: int = MAX_EXTRACT_BYTES,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> ExtractedDocument:
    """Extract plain text from trusted-size upload bytes, plus what the extractor saw.

    The result intentionally feeds the existing text/markdown chunker, so PDF/DOCX
    parsing does not alter the downstream embedding, ACL, or Qdrant payload path.
    """
    if not content:
        raise DocumentExtractionError("EMPTY_FILE", "Uploaded file is empty.")
    if len(content) > max_bytes:
        raise DocumentExtractionError(
            "FILE_TOO_LARGE",
            f"Uploaded file exceeds the {max_bytes} byte extraction limit.",
        )

    unit_count: int | None = None
    if mime_type in SUPPORTED_TEXT_MIME_TYPES:
        text = _decode_text(content)
        converter = TEXT_IDENTITY_CONVERTER
        unit_kind = UNIT_LINE
    elif mime_type == PDF_MIME_TYPE:
        text, unit_count = _extract_pdf_text(content, max_pdf_pages=max_pdf_pages)
        converter = f"pypdf/{_installed_version('pypdf')}"
        unit_kind = UNIT_PAGE
    elif mime_type == DOCX_MIME_TYPE:
        text, unit_count = _extract_docx_text(content)
        converter = f"python-docx/{_installed_version('python-docx')}"
        unit_kind = UNIT_PARAGRAPH
    else:
        raise DocumentExtractionError(
            "UNSUPPORTED_MIME_TYPE",
            f"Unsupported upload MIME type: {mime_type}",
        )

    normalized_text = text.replace("\r\n", "\n").replace("\r", "\n").strip()
    if not normalized_text:
        raise DocumentExtractionError(
            "EMPTY_EXTRACTED_TEXT",
            "Uploaded file did not contain extractable text.",
        )
    if unit_count is None:
        unit_count = len(normalized_text.split("\n"))
    return ExtractedDocument(
        text=normalized_text,
        converter=converter,
        source_unit_kind=unit_kind,
        source_unit_count=unit_count,
    )


def extract_text_from_bytes(
    *,
    mime_type: str,
    content: bytes,
    max_bytes: int = MAX_EXTRACT_BYTES,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> str:
    return extract_document(
        mime_type=mime_type,
        content=content,
        max_bytes=max_bytes,
        max_pdf_pages=max_pdf_pages,
    ).text


def extract_text(
    *,
    mime_type: str,
    file_bytes: bytes,
    max_bytes: int = MAX_EXTRACT_BYTES,
    max_pdf_pages: int = MAX_PDF_PAGES,
) -> str:
    return extract_text_from_bytes(
        mime_type=mime_type,
        content=file_bytes,
        max_bytes=max_bytes,
        max_pdf_pages=max_pdf_pages,
    )


def chunker_mime_type_for(mime_type: str) -> str:
    if mime_type in SUPPORTED_TEXT_MIME_TYPES:
        return mime_type
    if mime_type in SUPPORTED_BINARY_MIME_TYPES:
        return "text/plain"
    raise ValueError(f"Unsupported text parser MIME type: {mime_type}")


def _build_chunk(
    *,
    document_id: str,
    document_version: str,
    title: str,
    chunk_index: int,
    content: str,
    line_start: int,
    line_end: int,
    section_path: tuple[str, ...],
) -> ParsedChunk:
    content_hash = _sha256(content)
    locator = _citation_locator(title, section_path, line_start, line_end)
    locator_hash = hashlib.sha256(locator.encode("utf-8")).hexdigest()[:8]
    chunk_hash = _sha256(f"{content_hash}:{locator}")
    chunk_id = (
        f"{document_id}:{document_version}:l{line_start}-{line_end}:"
        f"c{chunk_index:03d}:{locator_hash}"
    )
    return ParsedChunk(
        chunk_id=chunk_id,
        chunk_index=chunk_index,
        content=content,
        content_hash=content_hash,
        chunk_hash=chunk_hash,
        token_count=len(content.split()),
        line_start=line_start,
        line_end=line_end,
        section_path=section_path,
        citation_locator=locator,
        parser_version=PARSER_VERSION,
        chunker_version=CHUNKER_VERSION,
    )


def _citation_locator(
    title: str,
    section_path: tuple[str, ...],
    line_start: int,
    line_end: int,
) -> str:
    section = " > ".join(section_path) if section_path else "body"
    return f"{title} / {section} / lines {line_start}-{line_end}"


def _parse_markdown_heading(line: str) -> tuple[int, str] | None:
    match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
    if match is None:
        return None
    return len(match.group(1)), match.group(2).strip()


def _sha256(value: str) -> str:
    return "sha256-" + hashlib.sha256(value.encode("utf-8")).hexdigest()


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="replace")


def _installed_version(distribution: str) -> str:
    """The installed version of an extractor library, or ``unknown``.

    Recorded so "which documents did the broken converter version touch" is answerable. It
    must never be able to fail an index attempt, hence the fallback.
    """
    try:
        return _package_version(distribution)
    except PackageNotFoundError:  # pragma: no cover - dependency is declared
        return "unknown"


def _extract_pdf_text(content: bytes, *, max_pdf_pages: int) -> tuple[str, int]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentExtractionError(
            "PARSER_DEPENDENCY_MISSING", "pypdf is required for PDF extraction."
        ) from exc

    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - untrusted file input
        raise DocumentExtractionError("PDF_PARSE_FAILED", str(exc)) from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 - untrusted file input
            raise DocumentExtractionError("PDF_ENCRYPTED", "Encrypted PDF cannot be parsed.") from exc
        if reader.is_encrypted:
            raise DocumentExtractionError("PDF_ENCRYPTED", "Encrypted PDF cannot be parsed.")

    if len(reader.pages) > max_pdf_pages:
        raise DocumentExtractionError(
            "PDF_PAGE_LIMIT_EXCEEDED",
            f"PDF has {len(reader.pages)} pages; limit is {max_pdf_pages}.",
        )

    page_texts: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - fail closed per page parse issue
            raise DocumentExtractionError(
                "PDF_PARSE_FAILED", f"Failed to extract text from page {page_number}: {exc}"
            ) from exc
        if page_text.strip():
            page_texts.append(page_text.strip())
    # The page COUNT is the document's, not the count of pages that yielded text: a scanned
    # page is a real page that contributed nothing, and that difference is the whole signal.
    return "\n\n".join(page_texts), len(reader.pages)


def _extract_docx_text(content: bytes) -> tuple[str, int]:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentExtractionError(
            "PARSER_DEPENDENCY_MISSING", "python-docx is required for DOCX extraction."
        ) from exc

    try:
        document = DocxDocument(BytesIO(content))
    except Exception as exc:  # noqa: BLE001 - untrusted file input
        raise DocumentExtractionError("DOCX_PARSE_FAILED", str(exc)) from exc

    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts), len(parts)

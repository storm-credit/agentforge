"""Which format a corpus document is ingested AS, and the fixture that requires.

WO-2026-08-14-EVAL-FORMAT-COVERAGE-001.

WHY THIS EXISTS

``live_runner`` used to register every corpus document with ``mime_type: text/markdown``
and a ``source_text`` body, so the corpus never touched the binary upload path that real
departments actually use. That is not a cosmetic difference: the product's
``chunker_mime_type_for`` maps PDF and DOCX to ``text/plain``, the markdown heading
detector therefore never runs on them, and the same policy text that yields several
clause-scoped chunks as markdown collapses into one section-path-less blob as DOCX. Every
number this harness has ever produced was measured on the one format that behaves best.

A corpus document may now declare ``ingestion_format``. ``markdown`` (the default, so every
existing corpus keeps its exact previous behaviour and its baselines stay comparable) takes
the unchanged register + index-job path. A binary format is uploaded as a real file through
``POST /knowledge/documents/upload``, i.e. the same endpoint an administrator uses, with a
fixture generated in memory.

THE FIXTURE IS NOT ALLOWED TO CHEAT

``build_docx_fixture`` renders a markdown body into a real Word document: ``## 제3장`` becomes
a paragraph carrying the ``Heading 2`` STYLE, not a paragraph whose text is literally
``## 제3장``. Writing the hash marks into the DOCX text would smuggle markdown through the
binary path, the heading detector would still not run (the chunker is handed ``text/plain``
either way), and the measurement would describe the fixture generator instead of the
product. The structure loss this harness reports has to be the product's real one.

PDF IS DELIBERATELY NOT SUPPORTED HERE (yet)

The product parses PDF (``pypdf``), but nothing in the dependency set WRITES one: ``pypdf``
is a reader and ``reportlab`` is not installed. A hand-rolled PDF containing Korean text
needs an embedded CJK font with a ``ToUnicode`` CMap, or ``pypdf`` extracts mojibake or
nothing at all -- so a PDF case would measure the quality of my improvised PDF writer, not
the product's ingestion. Adding PDF is one fixture builder plus one registry entry once a
PDF writer is an accepted dependency; ``fixture_for`` raises a named error until then rather
than silently ingesting a PDF case as something else.
"""

from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO

#: Registered, unchanged path: register the document as text/markdown + index-job source_text.
MARKDOWN = "markdown"
#: Real upload path: a generated .docx posted to POST /knowledge/documents/upload.
DOCX = "docx"
#: Declared by the product but not generatable here -- see the module docstring.
PDF = "pdf"

#: What a document with no ``ingestion_format`` gets. Every pre-existing corpus is markdown,
#: so the default keeps their requests, and therefore their baselines, byte-identical.
DEFAULT_INGESTION_FORMAT = MARKDOWN

SUPPORTED_INGESTION_FORMATS = (MARKDOWN, DOCX)
#: Formats that go through the real file-upload endpoint rather than the register path.
BINARY_INGESTION_FORMATS = (DOCX,)

MARKDOWN_MIME_TYPE = "text/markdown"
DOCX_MIME_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

#: format -> the MIME type the upload endpoint will see. Mirrors apps/api parsers.py; the
#: product remains the authority (the harness reads back what the product recorded).
UPLOAD_MIME_TYPES = {DOCX: DOCX_MIME_TYPE}
FILE_EXTENSIONS = {DOCX: ".docx"}

#: A case that is about no document at all (a refuse case with neither an expected nor a
#: forbidden document) and that declares no format of its own. Kept as its own bucket rather
#: than folded into markdown, because calling it markdown would be a claim nothing measured.
NO_DOCUMENT_FORMAT = "no_document"


class UnsupportedIngestionFormat(ValueError):
    """A corpus declared a format this harness cannot ingest honestly."""


@dataclass(frozen=True)
class UploadFixture:
    """The three things ``POST /knowledge/documents/upload`` needs for one file."""

    filename: str
    mime_type: str
    content: bytes


def normalize_format(value: str | None) -> str:
    """Resolve a declared ``ingestion_format`` (absent -> markdown), rejecting anything else.

    Fails loudly: a typo'd or not-yet-supported format must never be quietly ingested as
    markdown, because that is precisely the failure this Work Order exists to end.
    """
    if value is None:
        return DEFAULT_INGESTION_FORMAT
    normalized = str(value).strip().lower()
    if normalized not in SUPPORTED_INGESTION_FORMATS:
        supported = ", ".join(SUPPORTED_INGESTION_FORMATS)
        raise UnsupportedIngestionFormat(
            f"ingestion_format {value!r} is not supported by this harness (supported: "
            f"{supported}). PDF is parsed by the product but cannot be generated here -- "
            "see agentforge_eval/ingest_formats.py."
        )
    return normalized


def document_ingestion_format(document: dict) -> str:
    return normalize_format(document.get("ingestion_format"))


def document_formats(documents) -> dict[str, str]:
    """``doc_id -> declared format`` for a corpus's documents."""
    return {doc["doc_id"]: document_ingestion_format(doc) for doc in documents}


def case_ingestion_format(case: dict, doc_formats: dict[str, str]) -> str:
    """Which ingestion condition a case is measured under.

    An explicit ``ingestion_format`` on the case wins. Otherwise the format is inherited from
    the document the case is ABOUT -- the one it must cite, or, for a deny case, the one it
    must not leak. A case about no document at all falls back to ``NO_DOCUMENT_FORMAT``
    instead of being attributed to a format it never exercised.
    """
    declared = case.get("ingestion_format")
    if declared is not None:
        return normalize_format(declared)
    for key in ("expected_citation_doc", "forbidden_doc"):
        doc_id = case.get(key)
        if doc_id and doc_id in doc_formats:
            return doc_formats[doc_id]
    return NO_DOCUMENT_FORMAT


def is_binary_format(ingestion_format: str) -> bool:
    return ingestion_format in BINARY_INGESTION_FORMATS


def fixture_for(ingestion_format: str, *, doc_id: str, body: str) -> UploadFixture:
    """Generate the upload fixture for a declared binary format."""
    if ingestion_format == DOCX:
        return UploadFixture(
            filename=f"{doc_id}{FILE_EXTENSIONS[DOCX]}",
            mime_type=DOCX_MIME_TYPE,
            content=build_docx_fixture(body),
        )
    raise UnsupportedIngestionFormat(
        f"No upload fixture builder for ingestion_format {ingestion_format!r}."
    )


def build_docx_fixture(body: str) -> bytes:
    """Render a markdown body as a real .docx, headings as Word heading STYLES.

    ``# X`` -> Heading 1, ``## X`` -> Heading 2, everything else a body paragraph. The hash
    marks are consumed, never written into the text: see the module docstring.
    """
    from docx import Document as DocxDocument  # local import: only binary cases need it

    document = DocxDocument()
    for raw_line in body.replace("\r\n", "\n").replace("\r", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            continue
        level = len(line) - len(line.lstrip("#"))
        if level and line[level:].startswith(" "):
            document.add_heading(line[level:].strip(), level=min(level, 9))
        else:
            document.add_paragraph(line)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
